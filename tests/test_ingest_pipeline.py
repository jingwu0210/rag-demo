"""Task 4: Ingest Pipeline 测试

覆盖：
1. HierarchicalChunker（真实测试）：标题层级 / heading_path / heading_context / overlap / min 合并
2. DedupPipeline：MD5 精确去重 + duplicate_sources 记录
3. ConflictDetector：同 heading_path 不同文本 → 冲突
4. VersionedIngestService：真实 ChromaStore（tempfile）+ 假 embedding，ingested → replaced → skipped
5. OCRPipeline：仅清洗函数 + 原生 PDF 提取（fitz 内存生成），不跑 PaddleOCR 路径

说明：PaddleOCR（import 慢/需模型）与 BGE-M3 模型加载（需联网下载）不在测试范围，
embedder 用假向量 [0.1]*16 替代，详见 task-4-report.md。
"""
import os

import pytest

from core.config import ConfigRegistry
from core.ocr import ParsedDoc, OCRPipeline
from core.chunker import HierarchicalChunker, Chunk
from core.dedup import DedupPipeline, ConflictDetector
from core.versioned import VersionedIngestService
from storage.chroma_client import ChromaStore


# ── helpers ──────────────────────────────────────────────

def _body(seed: str, repeat: int = 17) -> str:
    """中文段落，约 100~500 token（中文 1 字符 ≈ 1 token）"""
    return seed * repeat


def _long_para(n: int = 22) -> str:
    """单个段落约 88 token（4 字符 × 22）"""
    return "会议要点" * n


def _chunk(cid: str, text: str, heading_path: str, source: str = "doc.pdf",
           language: str = "zh", embedding=None) -> Chunk:
    return Chunk(
        chunk_id=cid, text=text, heading_path=heading_path,
        language=language, embedding=embedding,
        metadata={"source_file": source, "heading_path": heading_path,
                  "language": language, "chunk_index": 0},
    )


# ── 1. chunker ───────────────────────────────────────────

def test_hierarchical_chunker_headings_and_context():
    ConfigRegistry.init("config.yaml")
    text = (
        "本手册适用于全体员工。\n\n"
        "# 员工手册\n"
        + _body("公司简介：本手册是全体员工的日常行为指南。") + "\n\n"
        "## 1.1 适用范围\n"
        + _body("本手册适用于公司全体正式员工及试用期员工。") + "\n\n"
        "## 第三章 薪酬福利\n"
        + _body("薪酬由基本工资、绩效奖金与福利津贴构成。") + "\n\n"
        "### 年假\n"
        + _body("员工入职满一年后每年享有年假。") + "\n\n"
        "第四章 考勤制度\n"
        + _body("每日上下班需打卡，迟到早退按规定处理。") + "\n\n"
        "## 4.1 打卡\n"
        + _body("打卡记录以系统时间为准。") + "\n\n"
        "§ 5 加班管理\n"
        + _body("加班需提前申请并经主管审批。") + "\n\n"
        "## 附录 A 会议记录\n"
        + "\n\n".join(_long_para() for _ in range(6)) + "\n\n"
        "### 附录 B 临时说明\n"
        + "后续补充事项。" + "\n"
    )

    doc = ParsedDoc(text=text, pages=[], tables=[], source="data/corpus/handbook.pdf",
                    language="zh")
    chunks = HierarchicalChunker().chunk(doc)

    # 基础断言：多 chunk
    assert len(chunks) > 1

    # 标题层级识别：heading_path 正确拼接（含 brief 示例路径）
    paths = {c.heading_path for c in chunks}
    assert "员工手册 > 第三章 薪酬福利 > 年假" in paths
    assert "员工手册 > 1.1 适用范围" in paths
    assert "第四章 考勤制度" in paths
    assert "第四章 考勤制度 > 4.1 打卡" in paths
    assert "第四章 考勤制度 > § 5 加班管理" in paths

    # 根内容（无标题）heading_path 为空
    assert any(c.heading_path == "" and "本手册适用于全体员工" in c.text for c in chunks)

    # heading_context 注入：有 heading_path 的 chunk 以 [ 开头，根 chunk 不注入
    for c in chunks:
        if c.heading_path:
            assert c.text.startswith("["), f"chunk {c.chunk_id} 缺少 heading_context"
        else:
            assert not c.text.startswith("[")

    # metadata 初始化：source 用文件名；chunk_index 连续
    assert chunks[0].metadata["source_file"] == "handbook.pdf"
    assert [c.metadata["chunk_index"] for c in chunks] == list(range(len(chunks)))


def test_hierarchical_chunker_long_section_overlap():
    ConfigRegistry.init("config.yaml")
    # 6 段 × 88 token = 528 > 512 → 触发段落切分 + overlap
    text = (
        "# 员工手册\n"
        + _body("员工手册正文内容。") + "\n\n"
        "## 附录 A 会议记录\n"
        + "\n\n".join(_long_para() for _ in range(6)) + "\n"
    )
    doc = ParsedDoc(text=text, pages=[], tables=[], source="handbook.pdf", language="zh")
    chunks = HierarchicalChunker().chunk(doc)

    section_chunks = [c for c in chunks if c.heading_path == "员工手册 > 附录 A 会议记录"]
    assert len(section_chunks) >= 2, "超长 section 应按段落切分成多个 chunk"

    c0, c1 = section_chunks[0], section_chunks[1]
    raw0 = c0.text.split("\n", 1)[1] if "\n" in c0.text else c0.text
    # overlap 携带上一 chunk 末尾 overlap_tokens 个 token（中文 64 字符）
    assert raw0[-64:] in c1.text, "chunk 切换时应携带上一 chunk 末尾的重叠前缀"


def test_hierarchical_chunker_min_merge():
    ConfigRegistry.init("config.yaml")
    text = (
        "# 员工手册\n"
        + _body("员工手册正文内容。") + "\n\n"
        "## 附录 C 简短条目\n"
        "一句话说明。" + "\n"
    )
    doc = ParsedDoc(text=text, pages=[], tables=[], source="handbook.pdf", language="zh")
    chunks = HierarchicalChunker().chunk(doc)

    # 附录 C（远小于 min_chunk_tokens）应并入上一 chunk，不产生独立 chunk
    assert not any(c.heading_path == "员工手册 > 附录 C 简短条目" for c in chunks)
    holder = [c for c in chunks if "[员工手册 > 附录 C 简短条目]" in c.text]
    assert len(holder) == 1
    assert holder[0].heading_path == "员工手册"
    assert "一句话说明" in holder[0].text


# ── 2. dedup ─────────────────────────────────────────────

def test_exact_dedup_keeps_one_and_records_sources():
    c1 = _chunk("a", "相同的政策文本", "员工手册 > 年假", source="doc1.pdf")
    c2 = _chunk("b", "相同的政策文本", "员工手册 > 年假", source="doc2.pdf")
    c3 = _chunk("c", "不同的政策文本", "员工手册 > 病假", source="doc3.pdf")

    result = DedupPipeline().exact_dedup([c1, c2, c3])

    assert len(result) == 2
    kept = [c for c in result if c.chunk_id == "a"][0]
    assert kept.metadata["duplicate_sources"] == ["doc2.pdf"]
    assert all(c.chunk_id != "b" for c in result)


# ── 3. conflict ──────────────────────────────────────────

def test_conflict_detector_same_heading_different_text():
    new = _chunk("n1", "新版年假政策：每年 15 天", "员工手册 > 年假", source="new.pdf")
    existing = [{
        "chunk_id": "e1",
        "text": "旧版年假政策：每年 5 天",
        "metadata": {"heading_path": "员工手册 > 年假", "source_file": "old.pdf",
                     "is_active": True},
    }]
    conflicts = ConflictDetector().detect([new], existing)

    assert len(conflicts) == 1
    c = conflicts[0]
    assert c["heading_path"] == "员工手册 > 年假"
    assert c["new_chunk_id"] == "n1"
    assert c["existing_chunk_id"] == "e1"
    assert c["new_source"] == "new.pdf"
    assert c["existing_source"] == "old.pdf"


def test_conflict_detector_same_text_no_conflict():
    new = _chunk("n1", "年假政策：每年 15 天", "员工手册 > 年假", source="new.pdf")
    existing = [{
        "chunk_id": "e1",
        "text": "年假政策：每年 15 天",
        "metadata": {"heading_path": "员工手册 > 年假", "source_file": "old.pdf",
                     "is_active": True},
    }]
    assert ConflictDetector().detect([new], existing) == []


# ── 4. versioned ingest ──────────────────────────────────

def test_versioned_ingest_flow(tmp_path):
    ConfigRegistry.init("config.yaml")
    ConfigRegistry.override("chromadb.persist_directory", str(tmp_path))
    store = ChromaStore()
    svc = VersionedIngestService(store)

    file_path = os.path.join(tmp_path, "handbook.txt")

    # 第一次提交（v1 内容）→ ingested
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("员工手册 v1 内容")
    h1 = svc.compute_hash(file_path)
    chunks1 = [_chunk("c1", "员工手册 v1 正文", "员工手册", source="handbook.txt",
                      embedding=[0.1] * 16)]
    res1 = svc.commit(chunks1, file_path, "handbook", "v1.0", h1)
    assert res1.status == "ingested"
    assert res1.chunks_created == 1 and res1.chunks_replaced == 0
    assert res1.source_file == "handbook.txt" and res1.version == "v1.0"
    assert svc.check_exists(h1) is True

    # 同 stem 不同内容（v2）→ replaced：旧版软下线、新版生效
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("员工手册 v2 更新内容")
    h2 = svc.compute_hash(file_path)
    assert h2 != h1
    chunks2 = [_chunk("c2", "员工手册 v2 正文", "员工手册", source="handbook.txt",
                      embedding=[0.1] * 16)]
    res2 = svc.commit(chunks2, file_path, "handbook", "v2.0", h2)
    assert res2.status == "replaced"
    assert res2.chunks_replaced == 1 and res2.chunks_created == 1

    # 相同 hash 再提交 → skipped
    res3 = svc.commit(chunks2, file_path, "handbook", "v2.0", h2)
    assert res3.status == "skipped"
    assert res3.reason != ""

    # 版本状态核对：c1 失效、c2 生效；metadata 字段完整
    items = store.collection.get(where={"source_file_stem": "handbook"})
    by_id = {m["chunk_id"]: m for m in items["metadatas"]}
    assert by_id["c1"]["is_active"] is False
    assert by_id["c2"]["is_active"] is True
    meta = by_id["c2"]
    for key in ("chunk_id", "source_file", "source_file_stem", "doc_type", "version",
                "effective_date", "ingested_at", "doc_hash", "language",
                "heading_path", "chunk_index", "is_active"):
        assert key in meta
    assert meta["doc_type"] == "handbook" and meta["version"] == "v2.0"
    assert meta["doc_hash"] == h2 and meta["source_file"] == "handbook.txt"

    # 查重语义：v1 已下线 → 不存在；v2 生效 → 存在
    assert svc.check_exists(h1) is False
    assert svc.check_exists(h2) is True


def test_versioned_commit_requires_embeddings(tmp_path):
    ConfigRegistry.init("config.yaml")
    ConfigRegistry.override("chromadb.persist_directory", str(tmp_path))
    store = ChromaStore()
    svc = VersionedIngestService(store)
    chunks = [_chunk("x1", "无向量的文本", "员工手册")]  # embedding=None
    with pytest.raises(ValueError):
        svc.commit(chunks, "handbook.txt", "handbook", "v1.0", "fakehash")


# ── 5. OCR（仅清洗 + 原生提取，不跑 PaddleOCR）──────────

def test_ocr_clean_normalize_whitespace():
    ocr = OCRPipeline()
    assert ocr._normalize_whitespace("a\n\n\n\nb") == "a\n\nb"
    assert ocr._normalize_whitespace("  a   b  ") == "a b"
    assert ocr._normalize_whitespace("line1\n\t line2") == "line1\n line2"


def test_ocr_clean_fix_cn_en_spacing():
    ocr = OCRPipeline()
    assert ocr._fix_cn_en_spacing("中文English混合") == "中文 English 混合"
    assert ocr._fix_cn_en_spacing("API接口规范") == "API 接口规范"
    assert ocr._fix_cn_en_spacing("纯中文不变") == "纯中文不变"


def test_ocr_clean_remove_header_footer():
    """I-3 终审：不再盲切首尾行 — 仅删页码行（纯数字/第N页/Page N）与空行，其余保留"""
    ocr = OCRPipeline()
    # 页眉"页眉公司名"非页码 → 保留；空行 → 删；"3"、"第4页页脚"、"Page 5" → 删
    text = "页眉公司名\n\n第二行正文\n第三行正文\n3\n第4页页脚\nPage 5"
    assert ocr._remove_header_footer(text) == "页眉公司名\n第二行正文\n第三行正文"
    # 无页码/空行 → 原样保留（不再盲切首尾行）
    assert ocr._remove_header_footer("a\nb") == "a\nb"
    assert ocr._remove_header_footer("首行内容\n正文\n末行内容") == "首行内容\n正文\n末行内容"


def test_ocr_native_pdf_extraction(tmp_path):
    import fitz

    ConfigRegistry.init("config.yaml")
    ConfigRegistry.override("ocr.engine", "")  # 关闭 OCR 路径，避免触发 PaddleOCR
    pdf_path = os.path.join(str(tmp_path), "native.pdf")
    doc = fitz.open()
    page = doc.new_page(width=120, height=120)
    for i in range(6):
        page.insert_text((10, 15 + i * 15), f"RAG line number {i}", fontsize=8)
    doc.save(pdf_path)
    doc.close()

    parsed = OCRPipeline().process(pdf_path)
    assert "RAG line number" in parsed.text
    assert len(parsed.pages) == 1
    assert parsed.pages[0]["num"] == 0
    assert parsed.source == pdf_path
    assert parsed.language == "en"

    ConfigRegistry.override("ocr.engine", "paddleocr")  # 恢复，避免污染其他测试


def test_ocr_markdown_file_parsing(tmp_path):
    """Phase 11: .md 文件走纯文本路径（不触发 fitz PDF 解析）"""
    md_path = os.path.join(str(tmp_path), "api_spec.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# API Specification\n\nRate limit: 1000 requests per minute.\n")

    parsed = OCRPipeline().process(md_path)
    assert "API Specification" in parsed.text
    assert "1000 requests" in parsed.text
    assert parsed.source == md_path
    assert parsed.language == "en"


def test_versioned_commit_custom_effective_date(tmp_path):
    """Phase 11: commit 支持显式 effective_date（过期文档埋点用）"""
    ConfigRegistry.init("config.yaml")
    ConfigRegistry.override("chromadb.persist_directory", str(tmp_path))
    store = ChromaStore()
    svc = VersionedIngestService(store)

    file_path = os.path.join(str(tmp_path), "legacy.txt")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("legacy manual")
    h = svc.compute_hash(file_path)
    chunks = [_chunk("c1", "legacy 正文", "遗留手册", source="legacy.txt",
                     embedding=[0.1] * 16)]
    res = svc.commit(chunks, file_path, "technical", "v1.0", h,
                     effective_date="2022-01-01")
    assert res.status == "ingested"

    items = store.collection.get(where={"chunk_id": "c1"})
    assert items["metadatas"][0]["effective_date"] == 20220101  # int YYYYMMDD


def test_versioned_doc_group_replaces_different_stems(tmp_path):
    """Phase 11 回归: 文件名带版本后缀（stem 不同）时靠 doc_group 替换旧版"""
    ConfigRegistry.init("config.yaml")
    ConfigRegistry.override("chromadb.persist_directory", str(tmp_path))
    store = ChromaStore()
    svc = VersionedIngestService(store)

    f1 = os.path.join(str(tmp_path), "handbook_v1.0.txt")
    with open(f1, "w", encoding="utf-8") as f:
        f.write("年假 5 天")
    h1 = svc.compute_hash(f1)
    res1 = svc.commit([_chunk("c1", "年假 5 天", "员工手册", source="handbook_v1.0.txt",
                              embedding=[0.1] * 16)],
                      f1, "handbook", "v1.0", h1, doc_group="employee_handbook")
    assert res1.status == "ingested"

    f2 = os.path.join(str(tmp_path), "handbook_v1.1.txt")
    with open(f2, "w", encoding="utf-8") as f:
        f.write("年假 10 天")
    h2 = svc.compute_hash(f2)
    res2 = svc.commit([_chunk("c2", "年假 10 天", "员工手册", source="handbook_v1.1.txt",
                              embedding=[0.1] * 16)],
                      f2, "handbook", "v1.1", h2, doc_group="employee_handbook")
    assert res2.status == "replaced"
    assert res2.chunks_replaced == 1

    items = store.collection.get(where={"doc_group": "employee_handbook"})
    by_id = {m["chunk_id"]: m for m in items["metadatas"]}
    assert by_id["c1"]["is_active"] is False   # 旧版软下线
    assert by_id["c2"]["is_active"] is True    # 新版生效
    assert by_id["c2"]["source_file_stem"] == "handbook_v1.1"  # stem 保留真实值
    assert by_id["c2"]["doc_group"] == "employee_handbook"


def test_ocr_text_page_not_misclassified_as_scanned(tmp_path):
    """Phase 11 回归: 正常文本页（~100 字符）不得误判为扫描件走 OCR 路径"""
    import fitz

    ConfigRegistry.init("config.yaml")
    pdf_path = os.path.join(str(tmp_path), "text_page.pdf")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((72, 72), "Normal text page with enough characters to be native text.",
                     fontsize=11)
    doc.save(pdf_path)
    doc.close()

    # 构造一个 OCR 路径会失败的引擎名：若走 OCR 会抛 ImportError，走原生则正常
    ConfigRegistry.override("ocr.engine", "nonexistent-engine")
    parsed = OCRPipeline().process(pdf_path)
    assert "Normal text page" in parsed.text
    ConfigRegistry.override("ocr.engine", "paddleocr")  # 恢复


def test_ocr_docx_parsing(tmp_path):
    """Phase 11: .docx 走 python-docx 提取路径（段落 + 表格）"""
    from docx import Document

    docx_path = os.path.join(str(tmp_path), "policy.docx")
    d = Document()
    d.add_paragraph("员工手册要点")
    d.add_paragraph("年假按照司龄计算。")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "职级"
    table.rows[0].cells[1].text = "天数"
    table.rows[1].cells[0].text = "P3"
    table.rows[1].cells[1].text = "10"
    d.save(docx_path)

    parsed = OCRPipeline().process(docx_path)
    assert "员工手册要点" in parsed.text
    assert "年假按照司龄计算" in parsed.text
    assert "| 职级 | 天数 |" in parsed.text   # 表格转文本
    assert parsed.source == docx_path
    assert parsed.language == "zh"

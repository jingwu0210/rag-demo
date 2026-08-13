import hashlib
import os
import re
from pathlib import Path

import fitz  # PyMuPDF
from core.config import ConfigRegistry
from core.bilingual import BilingualHandler

class ParsedDoc:
    def __init__(self, text: str, pages: list, tables: list, source: str, language: str):
        self.text = text
        self.pages = pages
        self.tables = tables
        self.source = source
        self.language = language

class OCRPipeline:
    """PDF 解析 + OCR。扫描件检测：逐页 text_ratio < 1% 视为扫描页，走 PaddleOCR。"""

    def process(self, file_path: str) -> ParsedDoc:
        ext = Path(file_path).suffix.lower()
        if ext in (".md", ".txt"):
            return self._process_text_file(file_path)
        if ext == ".docx":
            return self._process_docx(file_path)
        return self._process_pdf(file_path)

    def _process_docx(self, file_path: str) -> ParsedDoc:
        """Word 文档：python-docx 提取段落 + 表格，复用清洗逻辑"""
        from docx import Document
        clean_cfg = ConfigRegistry.get("ocr.clean", {})
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格按行转为 "| a | b |" 文本
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("| " + " | ".join(cells) + " |")
        text = "\n".join(parts)
        if clean_cfg.get("normalize_whitespace", True):
            text = self._normalize_whitespace(text)
        if clean_cfg.get("fix_cn_en_spacing", True):
            text = self._fix_cn_en_spacing(text)
        return ParsedDoc(text=text, pages=[{"num": 0, "text": text}], tables=[],
                         source=file_path, language=BilingualHandler.detect(text))

    def _process_text_file(self, file_path: str) -> ParsedDoc:
        """纯文本文件（.md/.txt）：直接读取 + 复用 PDF 路径的清洗逻辑"""
        clean_cfg = ConfigRegistry.get("ocr.clean", {})
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if clean_cfg.get("normalize_whitespace", True):
            text = self._normalize_whitespace(text)
        if clean_cfg.get("fix_cn_en_spacing", True):
            text = self._fix_cn_en_spacing(text)
        return ParsedDoc(text=text, pages=[{"num": 0, "text": text}], tables=[],
                         source=file_path, language=BilingualHandler.detect(text))

    def _process_pdf(self, file_path: str) -> ParsedDoc:
        doc = fitz.open(file_path)
        all_text, pages_data, tables_data = [], [], []
        need_ocr = ConfigRegistry.get("ocr.engine", "paddleocr")
        clean_cfg = ConfigRegistry.get("ocr.clean", {})
        # I-3 终审：OCR 临时图目录按文档区分（hash 前 8 位），并发 ingest 互不覆盖
        doc_tag = hashlib.md5(file_path.encode("utf-8")).hexdigest()[:8]

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            # 扫描件判定：按绝对字符数，不用字符数/像素面积比率。
            # 比率量纲错误（A4 页 50 万像素，正常文本页 500-1000 字符 → 比率恒 < 0.01，
            # 会把所有文本页误判为扫描件走 OCR，且 OCR 看不见白底白字的注入样本）
            char_count = len(text.strip())
            if char_count < 20 and need_ocr:
                text = self._ocr_page(page, page_num, doc_tag)
            else:
                text = self._extract_native(page)

            if clean_cfg.get("remove_header_footer", True):
                text = self._remove_header_footer(text)
            if clean_cfg.get("normalize_whitespace", True):
                text = self._normalize_whitespace(text)
            if clean_cfg.get("fix_cn_en_spacing", True):
                text = self._fix_cn_en_spacing(text)

            all_text.append(text)
            pages_data.append({"num": page_num, "text": text})

        full_text = "\n\n".join(all_text)
        language = BilingualHandler.detect(full_text)
        doc.close()
        return ParsedDoc(text=full_text, pages=pages_data, tables=tables_data, source=file_path, language=language)

    def _extract_native(self, page) -> str:
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (b[1], b[0]))
        return "\n".join(b[4] for b in blocks if b[4].strip())

    def _ocr_page(self, page, page_num: int, doc_tag: str) -> str:
        try:
            from paddleocr import PaddleOCR
        except ImportError as exc:
            raise RuntimeError(
                "PaddleOCR 未安装或版本不兼容（扫描件 OCR 不可用），"
                "请按 README 安装：pip install paddlepaddle paddleocr（macOS ARM "
                "如安装失败可尝试 paddlepaddle 2.x 组合）") from exc
        dpi = ConfigRegistry.get("ocr.dpi", 300)
        pix = page.get_pixmap(dpi=dpi)
        # I-3 终审：临时图路径按文档 hash 隔离（可追溯），用后 finally 清理
        ocr_cache = ConfigRegistry.get("paths.ocr_cache", "data/ocr")
        img_dir = os.path.join(ocr_cache, doc_tag)
        os.makedirs(img_dir, exist_ok=True)
        img_path = os.path.join(img_dir, f"page_{page_num}.png")
        pix.save(img_path)
        try:
            # I-3 终审：paddleocr 2.x/3.x 单模型单语言 — lang 必须是 str。
            # config ocr.language 为列表 → 取第一个元素；空列表默认 "ch"（中英混排）
            langs = ConfigRegistry.get("ocr.language", ["ch", "en"]) or []
            lang = langs[0] if langs else "ch"
            ocr = PaddleOCR(lang=lang, use_angle_cls=True)
            result = ocr.ocr(img_path, cls=True)
            if result and result[0]:
                return "\n".join(line[1][0] for line in result[0])
            return ""
        finally:
            try:
                os.remove(img_path)
            except OSError:
                pass
            try:
                os.rmdir(img_dir)
            except OSError:
                pass

    def _remove_header_footer(self, text: str) -> str:
        """I-3 终审：不再盲切首尾行。

        仅删除页码行（纯数字 / 第 N 页 / Page N）与完全空行；其余首尾行原样保留，
        避免误删正文首尾内容（如段首缩进行、末段未完行）。
        """
        page_patterns = (
            re.compile(r'^\s*\d+\s*$'),
            re.compile(r'^\s*第\s*\d+\s*页'),
            re.compile(r'^\s*Page\s*\d+', re.IGNORECASE),
        )
        kept = []
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if any(p.match(stripped) for p in page_patterns):
                continue
            kept.append(line)
        return "\n".join(kept)

    def _normalize_whitespace(self, text: str) -> str:
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        return text.strip()

    def _fix_cn_en_spacing(self, text: str) -> str:
        import re
        text = re.sub(r'([一-鿿])([a-zA-Z0-9])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z0-9])([一-鿿])', r'\1 \2', text)
        return text
